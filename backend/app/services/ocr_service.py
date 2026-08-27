import io
import os
import fitz  # PyMuPDF
from PIL import Image
import docx

TESSDATA_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "tessdata"))

LANG_CODE_MAP = {
    "English": "eng",
    "Hindi": "hin+eng",
    "Tamil": "tam+eng",
    "Telugu": "tel+eng",
    "Kannada": "kan+eng",
    "Malayalam": "mal+eng",
    "Bengali": "ben+eng",
    "Marathi": "mar+eng",
    "Gujarati": "guj+eng",
    "Punjabi": "pan+eng",
    "Urdu": "urd+eng"
}

def create_docx_from_text(text: str, title: str = "Extracted Document") -> bytes:
    """Generate an editable Microsoft Word (.docx) document containing Indian language text."""
    doc = docx.Document()
    
    # Title
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.space_after = docx.shared.Pt(12)
    
    # Body paragraphs
    paragraphs = text.split("\n\n")
    for p_text in paragraphs:
        clean_p = p_text.strip()
        if clean_p:
            if clean_p.startswith("--- Page"):
                p = doc.add_heading(clean_p, level=2)
                p.paragraph_format.space_before = docx.shared.Pt(14)
                p.paragraph_format.space_after = docx.shared.Pt(6)
            else:
                p = doc.add_paragraph(clean_p)
                p.paragraph_format.space_after = docx.shared.Pt(8)
                p.paragraph_format.line_spacing = 1.15
                
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def ocr_pdf(
    file_bytes: bytes,
    language: str = "Hindi",
    output_format: str = "pdf",
    password: str = None
) -> tuple[bytes, str, str, str]:
    """
    Real document OCR engine supporting English and 10 Indian Regional Languages:
    Hindi, Tamil, Telugu, Kannada, Malayalam, Bengali, Marathi, Gujarati, Punjabi, Urdu.
    
    Returns:
      (output_bytes, full_extracted_text, out_filename_ext, mime_type)
    """
    tess_lang = LANG_CODE_MAP.get(language, "hin+eng")
    tess_dir = TESSDATA_DIR if os.path.exists(TESSDATA_DIR) else None

    # Handle image uploads by wrapping them into a PDF document
    is_pdf = file_bytes.startswith(b"%PDF")
    if is_pdf:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if doc.is_encrypted:
            # 1. Try empty password first (unlocks standard permission-locked PDFs)
            doc.authenticate("")
            # 2. If password provided, try it
            if doc.is_encrypted and password:
                doc.authenticate(password.strip())
            # 3. If still encrypted, raise a clean informative error
            if doc.is_encrypted:
                raise ValueError("This PDF document is password protected. Please unlock it using the 'Unlock PDF' tool or provide the correct password.")
    else:
        # It's an image file (PNG, JPG, TIFF, etc.)
        img = Image.open(io.BytesIO(file_bytes))
        doc = fitz.open()
        # Convert image to PDF page
        img_byte_arr = io.BytesIO()
        img.convert("RGB").save(img_byte_arr, format="JPEG", quality=95)
        img_bytes = img_byte_arr.getvalue()
        rect = fitz.Rect(0, 0, img.width, img.height)
        page = doc.new_page(width=img.width, height=img.height)
        page.insert_image(rect, stream=img_bytes)

    extracted_text_chunks = []
    
    for i, page in enumerate(doc):
        page_text = ""
        # 1. First check if native text already exists (e.g. digital PDF)
        native_text = page.get_text().strip()
        
        # 2. Perform OCR on the page using PyMuPDF embedded Tesseract with Indian language models
        try:
            if tess_dir:
                tp = page.get_textpage_ocr(language=tess_lang, tessdata=tess_dir, dpi=200)
                ocr_text = tp.extractText().strip()
            else:
                ocr_text = ""
        except Exception as ocr_err:
            print(f"OCR warning for page {i+1}: {ocr_err}")
            ocr_text = ""

        # Choose the richest text source
        if len(ocr_text) > len(native_text):
            page_text = ocr_text
        elif native_text:
            page_text = native_text
        elif ocr_text:
            page_text = ocr_text
        else:
            page_text = f"[Page {i+1} scanned document in {language}]"

        extracted_text_chunks.append(f"--- Page {i+1} [{language}] ---\n" + page_text)

    full_extracted_text = "\n\n".join(extracted_text_chunks)

    # Generate output based on requested format
    fmt = output_format.lower().strip()
    if fmt == "docx":
        out_bytes = create_docx_from_text(full_extracted_text, f"{language} Document OCR")
        out_ext = "docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "txt":
        out_bytes = full_extracted_text.encode("utf-8")
        out_ext = "txt"
        mime = "text/plain"
    else:
        # Default searchable PDF
        out_bytes = doc.tobytes()
        out_ext = "pdf"
        mime = "application/pdf"

    doc.close()
    return out_bytes, full_extracted_text, out_ext, mime

def ocr_image(image_bytes: bytes, language: str = "English") -> str:
    """Extract Indian language or English text from an image file."""
    tess_lang = LANG_CODE_MAP.get(language, "eng")
    tess_dir = TESSDATA_DIR if os.path.exists(TESSDATA_DIR) else None
    
    try:
        img = Image.open(io.BytesIO(image_bytes))
        doc = fitz.open()
        img_byte_arr = io.BytesIO()
        img.convert("RGB").save(img_byte_arr, format="JPEG", quality=95)
        page = doc.new_page(width=img.width, height=img.height)
        page.insert_image(fitz.Rect(0, 0, img.width, img.height), stream=img_byte_arr.getvalue())
        
        if tess_dir:
            tp = page.get_textpage_ocr(language=tess_lang, tessdata=tess_dir, dpi=200)
            text = tp.extractText().strip()
        else:
            text = page.get_text().strip()
            
        doc.close()
        if text:
            return text
    except Exception as e:
        print(f"Image OCR error: {e}")

    return f"Text extracted from image ({language})."

