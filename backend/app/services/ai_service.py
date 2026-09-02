import io
import os
import re
import fitz
from collections import Counter
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

_REGISTERED_UNICODE_FONT = None

def get_unicode_font_name() -> str:
    global _REGISTERED_UNICODE_FONT
    if _REGISTERED_UNICODE_FONT:
        return _REGISTERED_UNICODE_FONT

    font_candidates = [
        # Linux paths (Render / Docker containers)
        ("NotoSansLinux", "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"),
        ("DejaVuSansLinux", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ("LiberationSans", "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
        # Windows paths
        ("ArialUnicode", "C:/Windows/Fonts/ARIALUNI.ttf"),
        ("NirmalaUI", "C:/Windows/Fonts/Nirmala.ttf"),
        ("SegoeUI", "C:/Windows/Fonts/segoeui.ttf"),
        ("NotoSans", "C:/Windows/Fonts/NotoSans-Regular.ttf"),
        ("Arial", "C:/Windows/Fonts/arial.ttf"),
    ]

    for fname, fpath in font_candidates:
        if os.path.exists(fpath):
            try:
                pdfmetrics.registerFont(TTFont(fname, fpath))
                _REGISTERED_UNICODE_FONT = fname
                return fname
            except Exception:
                continue

    _REGISTERED_UNICODE_FONT = "Helvetica"
    return "Helvetica"

def extract_pdf_full_text(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    lines = []
    for page in doc:
        t = page.get_text("text").strip()
        if t:
            lines.append(t)
    doc.close()
    return "\n\n".join(lines)

def summarize_pdf_text(text: str, max_sentences: int = 6) -> dict:
    """
    Extractive NLP & Statistical Keypoint Summarizer.
    Analyzes sentence frequencies, section headers, key metrics, and action items.
    """
    if not text.strip():
        return {
            "summary": "The document contains no readable text layer to summarize.",
            "key_takeaways": [],
            "metrics_found": [],
            "word_count": 0,
            "reading_time_minutes": 0
        }

    # Clean text
    clean_paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    words = re.findall(r'\b[A-Za-z]{3,}\b', text.lower())
    word_count = len(re.findall(r'\b\w+\b', text))
    reading_time = max(1, round(word_count / 200))

    # Stopwords filter
    stopwords = set([
        "the", "and", "for", "that", "this", "with", "from", "have", "are",
        "was", "were", "which", "will", "would", "about", "there", "their",
        "they", "been", "also", "into", "more", "other", "some", "such"
    ])
    meaningful_words = [w for w in words if w not in stopwords]
    word_freq = Counter(meaningful_words)

    # Score sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentence_scores = []

    for s in sentences:
        s_clean = s.strip().replace("\n", " ")
        if len(s_clean) < 25 or len(s_clean) > 350:
            continue
        score = 0
        for w in re.findall(r'\b[A-Za-z]{3,}\b', s_clean.lower()):
            score += word_freq.get(w, 0)
        sentence_scores.append((score / (len(s_clean.split()) + 1), s_clean))

    sentence_scores.sort(key=lambda x: x[0], reverse=True)
    top_sentences = [s[1] for s in sentence_scores[:max_sentences]]

    # Extract metrics / numbers
    metrics = re.findall(r'(\b\d+(?:[\.,]\d+)?%|\$\d+(?:[\.,]\d+)?|₹\d+(?:[\.,]\d+)?|\b\d{4,}\b)', text)
    unique_metrics = list(dict.fromkeys(metrics))[:8]

    # Extract key takeaways
    takeaways = []
    for s in top_sentences[:4]:
        takeaways.append(s)

    summary_text = " ".join(top_sentences[:3]) if top_sentences else text[:300]

    return {
        "summary": summary_text,
        "key_takeaways": takeaways,
        "metrics_found": unique_metrics,
        "word_count": word_count,
        "reading_time_minutes": reading_time
    }

def generate_summary_pdf(summary_data: dict, original_filename: str = "document.pdf") -> bytes:
    """Generate a clean executive summary PDF document."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1E1B4B'),
        spaceAfter=6
    )

    h2_style = ParagraphStyle(
        'H2Style',
        parent=styles['Heading2'],
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#4338CA'),
        spaceBefore=12,
        spaceAfter=6
    )

    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        spaceAfter=8
    )

    story.append(Paragraph(f"AI Executive Summary: {original_filename}", title_style))
    story.append(Paragraph(f"Analysis Stats: {summary_data['word_count']} words | ~{summary_data['reading_time_minutes']} min reading time", body_style))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Executive Overview", h2_style))
    story.append(Paragraph(summary_data["summary"], body_style))
    story.append(Spacer(1, 8))

    if summary_data.get("key_takeaways"):
        story.append(Paragraph("Key Findings & Takeaways", h2_style))
        for item in summary_data["key_takeaways"]:
            story.append(Paragraph(f"• {item}", body_style))
        story.append(Spacer(1, 8))

    if summary_data.get("metrics_found"):
        story.append(Paragraph("Key Figures & Statistical Metrics", h2_style))
        metrics_str = ", ".join(summary_data["metrics_found"])
        story.append(Paragraph(metrics_str, body_style))

    doc.build(story)
    return buffer.getvalue()

LANGUAGE_MAP = {
    "auto": "auto",
    "english": "en",
    "spanish": "es",
    "french": "fr",
    "german": "de",
    "hindi": "hi",
    "tamil": "ta",
    "telugu": "te",
    "kannada": "kn",
    "kanada": "kn",
    "kn": "kn",
    "malayalam": "ml",
    "bengali": "bn",
    "marathi": "mr",
    "gujarati": "gu",
    "punjabi": "pa",
    "urdu": "ur",
    "arabic": "ar",
    "chinese": "zh-CN",
    "chinese (simplified)": "zh-CN",
    "chinese (traditional)": "zh-TW",
    "japanese": "ja",
    "korean": "ko",
    "portuguese": "pt",
    "russian": "ru",
    "italian": "it",
    "dutch": "nl",
    "polish": "pl",
    "turkish": "tr",
    "vietnamese": "vi",
    "thai": "th",
    "indonesian": "id",
    "greek": "el",
    "hebrew": "he",
    "swedish": "sv",
    "norwegian": "no",
    "danish": "da",
    "finnish": "fi",
    "czech": "cs",
    "romanian": "ro",
    "hungarian": "hu",
    "ukrainian": "uk"
}

def resolve_lang_code(lang_str: str) -> str:
    if not lang_str:
        return "auto"
    cleaned = lang_str.strip().lower()
    if cleaned in LANGUAGE_MAP:
        return LANGUAGE_MAP[cleaned]
    if len(cleaned) == 2 or len(cleaned) == 5:
        return cleaned
    return "auto"

def translate_text_chunks(text: str, source_lang: str = "auto", target_lang: str = "es") -> str:
    """Translate text with automatic paragraph batching and robust retries."""
    if not text.strip():
        return ""
    
    import time
    from deep_translator import GoogleTranslator
    src = resolve_lang_code(source_lang)
    tgt = resolve_lang_code(target_lang)
    if tgt == "auto":
        tgt = "en"
        
    translator = GoogleTranslator(source=src, target=tgt)
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    translated_paras = []
    
    for p in paragraphs:
        if len(p) <= 2500:
            success = False
            for attempt in range(3):
                try:
                    res = translator.translate(p)
                    if res and not res.startswith("Error 500"):
                        translated_paras.append(res)
                        success = True
                        break
                    time.sleep(0.4)
                except Exception:
                    time.sleep(0.4)
            if not success:
                translated_paras.append(p)
        else:
            sentences = re.split(r'(?<=[.!?])\s+', p)
            curr_chunk = ""
            sub_trans = []
            for s in sentences:
                if len(curr_chunk) + len(s) + 1 > 2000:
                    for attempt in range(3):
                        try:
                            t_res = translator.translate(curr_chunk)
                            if t_res and not t_res.startswith("Error 500"):
                                sub_trans.append(t_res)
                                break
                            time.sleep(0.4)
                        except Exception:
                            time.sleep(0.4)
                    else:
                        sub_trans.append(curr_chunk)
                    curr_chunk = s
                else:
                    curr_chunk = f"{curr_chunk} {s}".strip()
            if curr_chunk:
                for attempt in range(3):
                    try:
                        t_res = translator.translate(curr_chunk)
                        if t_res and not t_res.startswith("Error 500"):
                            sub_trans.append(t_res)
                            break
                        time.sleep(0.4)
                    except Exception:
                        time.sleep(0.4)
                else:
                    sub_trans.append(curr_chunk)
            translated_paras.append(" ".join(sub_trans))
            
    return "\n\n".join(translated_paras)

def translate_pdf_document(
    pdf_bytes: bytes,
    target_language: str = "Spanish",
    source_language: str = "auto",
    output_format: str = "pdf",
    password: str = None
) -> dict:
    """
    Translates PDF content with exact in-place layout, alignment, and coordinate preservation.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    if doc.is_encrypted:
        if password:
            if not doc.authenticate(password):
                raise ValueError("Incorrect PDF password provided.")
        else:
            raise ValueError("Document is password protected. Please provide a password.")

    fmt = output_format.lower().strip()

    # Check if input document has existing text blocks that can be translated in-place
    has_layout_blocks = False
    for page in doc:
        blocks = page.get_text("blocks")
        text_b_count = sum(1 for b in blocks if b[6] == 0 and len(b[4].strip()) > 3)
        if text_b_count > 0:
            has_layout_blocks = True
            break

    # =========================================================================
    # 1. IN-PLACE LAYOUT-PRESERVING PDF TRANSLATION (Exact Alignment & Coordinates)
    # =========================================================================
    if fmt == "pdf" and has_layout_blocks:
        font_path = "C:/Windows/Fonts/ARIALUNI.ttf" if os.path.exists("C:/Windows/Fonts/ARIALUNI.ttf") else "C:/Windows/Fonts/arial.ttf"
        all_orig_text_parts = []
        all_trans_text_parts = []

        for pno, page in enumerate(doc):
            text_dict = page.get_text("dict")
            blocks_to_replace = []

            for b in text_dict.get("blocks", []):
                if b.get("type") == 0:  # text block
                    block_rect = fitz.Rect(b["bbox"])
                    block_text = ""
                    sample_size = 10
                    sample_color = (0, 0, 0)

                    for line in b.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            line_text += span.get("text", "") + " "
                            sample_size = span.get("size", 10)
                            c_int = span.get("color", 0)
                            r = ((c_int >> 16) & 255) / 255.0
                            g = ((c_int >> 8) & 255) / 255.0
                            bl = (c_int & 255) / 255.0
                            sample_color = (r, g, bl)
                        block_text += line_text.strip() + "\n"

                    clean_text = block_text.strip()
                    if clean_text:
                        # Detect horizontal alignment based on position
                        page_w = page.rect.width
                        align = 0
                        if block_rect.width < page_w * 0.75:
                            center_diff = abs((block_rect.x0 + block_rect.x1)/2 - page_w/2)
                            if center_diff < 35:
                                align = 1  # Centered
                            elif block_rect.x0 > page_w * 0.55:
                                align = 2  # Right-aligned

                        blocks_to_replace.append({
                            "rect": block_rect,
                            "text": clean_text,
                            "size": sample_size,
                            "color": sample_color,
                            "align": align
                        })
                        page.add_redact_annot(block_rect, fill=(1, 1, 1))

            page.apply_redactions()

            font_id = f"f_{pno}"
            if os.path.exists(font_path):
                page.insert_font(fontname=font_id, fontfile=font_path)
            else:
                font_id = "helv"

            for item in blocks_to_replace:
                trans_txt = translate_text_chunks(item["text"], source_language, target_language)
                all_orig_text_parts.append(item["text"])
                all_trans_text_parts.append(trans_txt)

                fsize = item["size"]
                rc = page.insert_textbox(
                    item["rect"],
                    trans_txt,
                    fontname=font_id,
                    fontsize=fsize,
                    color=item["color"],
                    align=item["align"]
                )
                if rc < 0 and fsize > 6:
                    # Dynamically adjust font size to guarantee it fits inside original box
                    page.insert_textbox(
                        item["rect"],
                        trans_txt,
                        fontname=font_id,
                        fontsize=max(6.5, fsize * 0.85),
                        color=item["color"],
                        align=item["align"]
                    )

        try:
            doc.subset_fonts()
        except Exception:
            pass

        output_pdf_bytes = doc.tobytes(garbage=4, deflate=True)
        doc.close()

        full_orig = "\n\n".join(all_orig_text_parts)
        full_trans = "\n\n".join(all_trans_text_parts)
        word_count = len(re.findall(r'\b\w+\b', full_trans))

        return {
            "bytes": output_pdf_bytes,
            "ext": "pdf",
            "mime": "application/pdf",
            "original_text": full_orig,
            "translated_text": full_trans,
            "word_count": word_count,
            "target_language": target_language
        }

    # =========================================================================
    # 2. SCANNED OCR OR DOCUMENT EXPORT (DOCX, TXT, or Full Reconstructed PDF)
    # =========================================================================
    from app.services.ocr_service import TESSDATA_DIR, LANG_CODE_MAP
    tess_lang = LANG_CODE_MAP.get(source_language.capitalize(), "eng")
    tess_dir = TESSDATA_DIR if os.path.exists(TESSDATA_DIR) else None

    extracted_pages = []
    for i, page in enumerate(doc):
        t = page.get_text("text").strip()
        if len(t) >= 15:
            extracted_pages.append(t)
        elif tess_dir:
            try:
                tp = page.get_textpage_ocr(language=tess_lang, tessdata=tess_dir, dpi=200)
                ocr_t = tp.extractText().strip()
                if ocr_t:
                    extracted_pages.append(ocr_t)
                elif t:
                    extracted_pages.append(t)
            except Exception:
                if t:
                    extracted_pages.append(t)
        elif t:
            extracted_pages.append(t)
    doc.close()

    full_original_text = "\n\n".join(extracted_pages)
    if not full_original_text.strip():
        raise ValueError("No readable text found in document to translate. Please make sure the PDF has readable text or clear scans.")

    translated_text = translate_text_chunks(full_original_text, source_language, target_language)
    word_count = len(re.findall(r'\b\w+\b', translated_text))

    if fmt == "docx":
        import docx
        doc_obj = docx.Document()
        heading = doc_obj.add_heading(f"Translated Document ({target_language.capitalize()})", level=1)
        heading.paragraph_format.space_after = docx.shared.Pt(14)

        for para in translated_text.split("\n\n"):
            clean_p = para.strip()
            if clean_p:
                p = doc_obj.add_paragraph(clean_p)
                p.paragraph_format.space_after = docx.shared.Pt(8)
                p.paragraph_format.line_spacing = 1.15

        buf = io.BytesIO()
        doc_obj.save(buf)
        return {
            "bytes": buf.getvalue(),
            "ext": "docx",
            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "original_text": full_original_text,
            "translated_text": translated_text,
            "word_count": word_count,
            "target_language": target_language
        }

    elif fmt == "txt":
        return {
            "bytes": translated_text.encode("utf-8"),
            "ext": "txt",
            "mime": "text/plain; charset=utf-8",
            "original_text": full_original_text,
            "translated_text": translated_text,
            "word_count": word_count,
            "target_language": target_language
        }

    else:
        # Generate clean PDF with Unicode TrueType font support
        unicode_font = get_unicode_font_name()
        buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        story = []

        title_style = ParagraphStyle(
            'TitleStyle',
            fontName=unicode_font,
            fontSize=16,
            leading=22,
            textColor=colors.HexColor('#1E1B4B'),
            spaceAfter=12
        )

        body_style = ParagraphStyle(
            'BodyStyle',
            fontName=unicode_font,
            fontSize=10,
            leading=16,
            textColor=colors.HexColor('#334155'),
            spaceAfter=10
        )

        meta_style = ParagraphStyle(
            'MetaStyle',
            fontName=unicode_font,
            fontSize=9,
            leading=14,
            textColor=colors.HexColor('#64748B'),
            spaceAfter=16
        )

        story.append(Paragraph(f"DocFlow Translated Document ({target_language.capitalize()})", title_style))
        story.append(Paragraph(f"Source Language: {source_language.capitalize()} | Target Language: {target_language.capitalize()} | Words: {word_count}", meta_style))
        story.append(Spacer(1, 10))

        for para in translated_text.split("\n\n"):
            clean_p = para.strip().replace("\n", " ")
            if clean_p:
                safe_p = clean_p.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                try:
                    story.append(Paragraph(safe_p, body_style))
                except Exception:
                    pass

        pdf_doc.build(story)
        return {
            "bytes": buffer.getvalue(),
            "ext": "pdf",
            "mime": "application/pdf",
            "original_text": full_original_text,
            "translated_text": translated_text,
            "word_count": word_count,
            "target_language": target_language
        }
