import io
import os
import re
import time
import uuid
import fitz  # PyMuPDF
import docx
from typing import Dict, Any, List, Optional, Tuple
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from deep_translator import GoogleTranslator
from langdetect import detect_langs

# Global in-memory job status store
_JOBS: Dict[str, Dict[str, Any]] = {}
_REGISTERED_UNICODE_FONT: Optional[str] = None

# Comprehensive Language Code Mapping
LANGUAGE_MAP: Dict[str, str] = {
    "auto": "auto",
    "english": "en",
    "hindi": "hi",
    "tamil": "ta",
    "telugu": "te",
    "kannada": "kn",
    "kanada": "kn",
    "malayalam": "ml",
    "bengali": "bn",
    "marathi": "mr",
    "gujarati": "gu",
    "punjabi": "pa",
    "urdu": "ur",
    "spanish": "es",
    "french": "fr",
    "german": "de",
    "portuguese": "pt",
    "italian": "it",
    "chinese": "zh-CN",
    "chinese (simplified)": "zh-CN",
    "chinese (traditional)": "zh-TW",
    "japanese": "ja",
    "korean": "ko",
    "arabic": "ar",
    "russian": "ru",
    "dutch": "nl",
    "polish": "pl",
    "turkish": "tr",
    "vietnamese": "vi",
    "thai": "th",
    "indonesian": "id",
    "greek": "el",
    "hebrew": "he",
    "swedish": "sv",
    "norwegian": "no",
    "danish": "da",
    "finnish": "fi",
    "czech": "cs",
    "romanian": "ro",
    "hungarian": "hu",
    "ukrainian": "uk"
}

LANGUAGE_NAMES: Dict[str, str] = {
    "en": "English",
    "hi": "Hindi",
    "ta": "Tamil",
    "te": "Telugu",
    "kn": "Kannada",
    "ml": "Malayalam",
    "bn": "Bengali",
    "mr": "Marathi",
    "gu": "Gujarati",
    "pa": "Punjabi",
    "ur": "Urdu",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "pt": "Portuguese",
    "it": "Italian",
    "zh-cn": "Chinese",
    "ja": "Japanese",
    "ko": "Korean",
    "ar": "Arabic",
    "ru": "Russian",
    "nl": "Dutch",
    "pl": "Polish",
    "tr": "Turkish",
    "vi": "Vietnamese",
    "th": "Thai",
    "id": "Indonesian",
    "el": "Greek",
    "he": "Hebrew",
    "sv": "Swedish",
    "no": "Norwegian",
    "da": "Danish",
    "fi": "Finnish",
    "cs": "Czech",
    "ro": "Romanian",
    "hu": "Hungarian",
    "uk": "Ukrainian"
}

def resolve_lang_code(lang_str: str) -> str:
    """Normalize language string to ISO 639-1 code."""
    if not lang_str:
        return "auto"
    cleaned = lang_str.strip().lower()
    if cleaned in LANGUAGE_MAP:
        return LANGUAGE_MAP[cleaned]
    if len(cleaned) in (2, 5):
        return cleaned
    return "auto"

def get_language_display_name(code: str) -> str:
    """Get readable language name for ISO code."""
    c = code.lower()
    return LANGUAGE_NAMES.get(c, c.upper())

def get_unicode_font_name() -> str:
    """Register and return a TrueType Unicode font supporting Indian and world scripts."""
    global _REGISTERED_UNICODE_FONT
    if _REGISTERED_UNICODE_FONT:
        return _REGISTERED_UNICODE_FONT

    font_candidates = [
        ("ArialUnicode", "C:/Windows/Fonts/ARIALUNI.ttf"),
        ("NirmalaUI", "C:/Windows/Fonts/Nirmala.ttf"),
        ("SegoeUI", "C:/Windows/Fonts/segoeui.ttf"),
        ("NotoSans", "C:/Windows/Fonts/NotoSans-Regular.ttf"),
        ("Arial", "C:/Windows/Fonts/arial.ttf"),
    ]

    for fname, fpath in font_candidates:
        if os.path.exists(fpath):
            try:
                pdfmetrics.registerFont(TTFont(fname, fpath))
                _REGISTERED_UNICODE_FONT = fname
                return fname
            except Exception:
                continue

    _REGISTERED_UNICODE_FONT = "Helvetica"
    return "Helvetica"

def create_translation_job() -> str:
    job_id = str(uuid.uuid4())
    _JOBS[job_id] = {
        "job_id": job_id,
        "status": "queued",
        "stage": "queued",
        "stage_label": "Queued for translation...",
        "current_page": 0,
        "total_pages": 0,
        "detected_language": "Unknown",
        "confidence": "Normal",
        "download_key": None,
        "filename": None,
        "error": None,
        "result": None,
        "created_at": time.time()
    }
    return job_id

def update_job_status(job_id: str, **kwargs):
    if job_id in _JOBS:
        _JOBS[job_id].update(kwargs)

def get_job_status(job_id: str) -> Optional[Dict[str, Any]]:
    return _JOBS.get(job_id)

def analyze_document_content(pdf_bytes: bytes, password: Optional[str] = None) -> Dict[str, Any]:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    if doc.is_encrypted:
        if password:
            if not doc.authenticate(password):
                raise ValueError("Incorrect PDF password provided.")
        else:
            raise ValueError("Document is password protected. Please provide a password.")

    total_pages = len(doc)
    page_texts = []
    scanned_pages = []
    total_words = 0

    for pno, page in enumerate(doc):
        t = page.get_text("text").strip()
        if len(t) < 20:
            scanned_pages.append(pno + 1)
        else:
            page_texts.append(t)
            total_words += len(t.split())

    full_sample_text = " ".join(page_texts)[:4000]
    detected_lang = "English"
    confidence = "High"
    is_multi_lang = False

    if full_sample_text.strip():
        try:
            detections = detect_langs(full_sample_text)
            if detections:
                top = detections[0]
                detected_lang = get_language_display_name(top.lang)
                if top.prob >= 0.85:
                    confidence = "High"
                elif top.prob >= 0.50:
                    confidence = "Medium"
                else:
                    confidence = "Low"

                if len(detections) > 1 and detections[1].prob > 0.30:
                    is_multi_lang = True
        except Exception:
            detected_lang = "English"
            confidence = "Medium"

    doc.close()

    return {
        "total_pages": total_pages,
        "scanned_pages": scanned_pages,
        "is_scanned": len(scanned_pages) == total_pages,
        "has_scanned_pages": len(scanned_pages) > 0,
        "detected_language": detected_lang,
        "confidence": confidence,
        "is_multi_language": is_multi_lang,
        "total_words": total_words
    }

def protect_tokens(text: str) -> Tuple[str, Dict[str, str]]:
    replacements = {}
    counter = 0

    def replace_match(match):
        nonlocal counter
        token = f"__TOKEN_{counter}__"
        replacements[token] = match.group(0)
        counter += 1
        return token

    patterns = [
        r'https?://[^\s<>"]+|www\.[^\s<>"]+',
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        r'₹\s*[\d,]+(?:\.\d+)?|\$\s*[\d,]+(?:\.\d+)?|€\s*[\d,]+(?:\.\d+)?|£\s*[\d,]+(?:\.\d+)?',
        r'\b\d+(?:\.\d+)?%',
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
        r'\b\+?\d{1,3}[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}\b'
    ]

    protected_text = text
    for pattern in patterns:
        protected_text = re.sub(pattern, replace_match, protected_text)

    return protected_text, replacements

def restore_tokens(text: str, replacements: Dict[str, str]) -> str:
    restored = text
    for token, original in replacements.items():
        token_clean = token.strip()
        restored = restored.replace(token_clean, original)
        restored = re.sub(r'__\s*TOKEN\s*_\s*(\d+)\s*__', lambda m: replacements.get(f"__TOKEN_{m.group(1)}__", m.group(0)), restored)
    return restored

class TranslationEngine:
    def __init__(self, source_lang: str = "auto", target_lang: str = "Spanish"):
        self.src_code = resolve_lang_code(source_lang)
        self.tgt_code = resolve_lang_code(target_lang)
        if self.tgt_code == "auto":
            self.tgt_code = "en"
        self.translator = GoogleTranslator(source=self.src_code, target=self.tgt_code)

    def translate_text(self, text: str) -> str:
        if not text.strip():
            return ""

        if self.src_code != "auto" and self.src_code == self.tgt_code:
            return text

        protected_text, replacements = protect_tokens(text)
        paragraphs = [p.strip() for p in protected_text.split("\n\n") if p.strip()]
        translated_paras = []

        for p in paragraphs:
            if len(p) <= 2500:
                trans = self._translate_chunk_with_retry(p)
                translated_paras.append(trans)
            else:
                sentences = re.split(r'(?<=[.!?])\s+', p)
                curr_chunk = ""
                sub_trans = []
                for s in sentences:
                    if len(curr_chunk) + len(s) + 1 > 2000:
                        sub_trans.append(self._translate_chunk_with_retry(curr_chunk))
                        curr_chunk = s
                    else:
                        curr_chunk = f"{curr_chunk} {s}".strip()
                if curr_chunk:
                    sub_trans.append(self._translate_chunk_with_retry(curr_chunk))
                translated_paras.append(" ".join(sub_trans))

        full_trans = "\n\n".join(translated_paras)
        return restore_tokens(full_trans, replacements)

    def _translate_chunk_with_retry(self, chunk: str, max_retries: int = 3) -> str:
        if not chunk.strip():
            return ""

        for attempt in range(max_retries):
            try:
                res = self.translator.translate(chunk)
                if res and not res.startswith("Error 500"):
                    return res
                time.sleep(0.4 * (attempt + 1))
            except Exception:
                time.sleep(0.4 * (attempt + 1))

        return chunk

def validate_output_pdf(pdf_bytes: bytes, min_pages: int = 1):
    if not pdf_bytes or len(pdf_bytes) < 100:
        raise ValueError("Generated PDF output is empty or corrupted.")

    try:
        check_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if len(check_doc) < min_pages:
            raise ValueError(f"Output PDF page count mismatch: expected at least {min_pages} pages, found {len(check_doc)}.")

        total_text_len = sum(len(page.get_text("text").strip()) for page in check_doc)
        check_doc.close()

        if total_text_len == 0:
            raise ValueError("Output PDF contains no readable text content.")
    except Exception as e:
        raise ValueError(f"Output PDF validation failed: {str(e)}")

def process_pdf_translation(
    pdf_bytes: bytes,
    target_language: str = "Spanish",
    source_language: str = "auto",
    output_format: str = "pdf",
    password: Optional[str] = None,
    job_id: Optional[str] = None
) -> Dict[str, Any]:
    if job_id:
        update_job_status(job_id, status="processing", stage="analyzing", stage_label="Analyzing document structure...")

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    if doc.is_encrypted:
        if password:
            if not doc.authenticate(password):
                raise ValueError("Incorrect PDF password provided.")
        else:
            raise ValueError("Document is password protected. Please provide a password.")

    total_pages = len(doc)
    if job_id:
        update_job_status(job_id, total_pages=total_pages)

    analysis = analyze_document_content(pdf_bytes, password)
    detected_lang = analysis["detected_language"]
    confidence = analysis["confidence"]

    if job_id:
        update_job_status(
            job_id,
            stage="detecting",
            stage_label=f"Detected language: {detected_lang} (Confidence: {confidence})",
            detected_language=detected_lang,
            confidence=confidence
        )

    effective_source = detected_lang if source_language == "auto" else source_language
    engine = TranslationEngine(source_lang=effective_source, target_lang=target_language)

    fmt = output_format.lower().strip()

    has_layout_blocks = False
    for page in doc:
        blocks = page.get_text("blocks")
        text_b_count = sum(1 for b in blocks if b[6] == 0 and len(b[4].strip()) > 3)
        if text_b_count > 0:
            has_layout_blocks = True
            break

    font_path = "C:/Windows/Fonts/ARIALUNI.ttf" if os.path.exists("C:/Windows/Fonts/ARIALUNI.ttf") else "C:/Windows/Fonts/arial.ttf"

    if fmt == "pdf" and has_layout_blocks:
        all_orig_text_parts = []
        all_trans_text_parts = []

        for pno, page in enumerate(doc):
            if job_id:
                update_job_status(
                    job_id,
                    stage="translating",
                    stage_label=f"Translating page {pno + 1} of {total_pages}...",
                    current_page=pno + 1
                )

            text_dict = page.get_text("dict")
            blocks_to_replace = []

            for b in text_dict.get("blocks", []):
                if b.get("type") == 0:
                    block_rect = fitz.Rect(b["bbox"])
                    block_text = ""
                    sample_size = 10
                    sample_color = (0, 0, 0)

                    for line in b.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            line_text += span.get("text", "") + " "
                            sample_size = span.get("size", 10)
                            c_int = span.get("color", 0)
                            r = ((c_int >> 16) & 255) / 255.0
                            g = ((c_int >> 8) & 255) / 255.0
                            bl = (c_int & 255) / 255.0
                            sample_color = (r, g, bl)
                        block_text += line_text.strip() + "\n"

                    clean_text = block_text.strip()
                    if clean_text:
                        page_w = page.rect.width
                        align = 0
                        if block_rect.width < page_w * 0.75:
                            center_diff = abs((block_rect.x0 + block_rect.x1)/2 - page_w/2)
                            if center_diff < 35:
                                align = 1
                            elif block_rect.x0 > page_w * 0.55:
                                align = 2

                        blocks_to_replace.append({
                            "rect": block_rect,
                            "text": clean_text,
                            "size": sample_size,
                            "color": sample_color,
                            "align": align
                        })
                        page.add_redact_annot(block_rect, fill=(1, 1, 1))

            page.apply_redactions()

            font_id = f"f_{pno}"
            if os.path.exists(font_path):
                page.insert_font(fontname=font_id, fontfile=font_path)
            else:
                font_id = "helv"

            for item in blocks_to_replace:
                trans_txt = engine.translate_text(item["text"])
                all_orig_text_parts.append(item["text"])
                all_trans_text_parts.append(trans_txt)

                fsize = item["size"]
                rc = page.insert_textbox(
                    item["rect"],
                    trans_txt,
                    fontname=font_id,
                    fontsize=fsize,
                    color=item["color"],
                    align=item["align"]
                )
                if rc < 0 and fsize > 6:
                    page.insert_textbox(
                        item["rect"],
                        trans_txt,
                        fontname=font_id,
                        fontsize=max(6.5, fsize * 0.85),
                        color=item["color"],
                        align=item["align"]
                    )

        if job_id:
            update_job_status(job_id, stage="generating", stage_label="Reconstructing layout-preserved PDF...")

        try:
            doc.subset_fonts()
        except Exception:
            pass

        output_pdf_bytes = doc.tobytes(garbage=4, deflate=True)
        doc.close()

        if job_id:
            update_job_status(job_id, stage="validating", stage_label="Validating output quality...")
        validate_output_pdf(output_pdf_bytes, min_pages=total_pages)

        full_orig = "\n\n".join(all_orig_text_parts)
        full_trans = "\n\n".join(all_trans_text_parts)
        word_count = len(re.findall(r'\b\w+\b', full_trans))

        return {
            "bytes": output_pdf_bytes,
            "ext": "pdf",
            "mime": "application/pdf",
            "original_text": full_orig,
            "translated_text": full_trans,
            "word_count": word_count,
            "detected_language": detected_lang,
            "confidence": confidence,
            "target_language": target_language
        }

    from app.services.ocr_service import TESSDATA_DIR, LANG_CODE_MAP
    tess_lang = LANG_CODE_MAP.get(effective_source.capitalize(), "eng")
    tess_dir = TESSDATA_DIR if os.path.exists(TESSDATA_DIR) else None

    extracted_pages = []
    for pno, page in enumerate(doc):
        if job_id:
            update_job_status(
                job_id,
                stage="ocr" if len(page.get_text().strip()) < 15 else "extracting",
                stage_label=f"Extracting content on page {pno + 1} of {total_pages}...",
                current_page=pno + 1
            )

        t = page.get_text("text").strip()
        if len(t) >= 15:
            extracted_pages.append(t)
        elif tess_dir:
            try:
                tp = page.get_textpage_ocr(language=tess_lang, tessdata=tess_dir, dpi=200)
                ocr_t = tp.extractText().strip()
                if ocr_t:
                    extracted_pages.append(ocr_t)
                elif t:
                    extracted_pages.append(t)
            except Exception:
                if t:
                    extracted_pages.append(t)
        elif t:
            extracted_pages.append(t)
    doc.close()

    full_original_text = "\n\n".join(extracted_pages)
    if not full_original_text.strip():
        raise ValueError("No readable text found in document to translate. Please ensure the document has clear readable text or clear scans.")

    if job_id:
        update_job_status(job_id, stage="translating", stage_label="Translating document content...")

    translated_text = engine.translate_text(full_original_text)
    word_count = len(re.findall(r'\b\w+\b', translated_text))

    if fmt == "docx":
        if job_id:
            update_job_status(job_id, stage="generating", stage_label="Generating editable Word (.docx) document...")

        doc_obj = docx.Document()
        heading = doc_obj.add_heading(f"Translated Document ({target_language.capitalize()})", level=1)
        heading.paragraph_format.space_after = docx.shared.Pt(14)

        for para in translated_text.split("\n\n"):
            clean_p = para.strip()
            if clean_p:
                p = doc_obj.add_paragraph(clean_p)
                p.paragraph_format.space_after = docx.shared.Pt(8)
                p.paragraph_format.line_spacing = 1.15

        buf = io.BytesIO()
        doc_obj.save(buf)
        return {
            "bytes": buf.getvalue(),
            "ext": "docx",
            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "original_text": full_original_text,
            "translated_text": translated_text,
            "word_count": word_count,
            "detected_language": detected_lang,
            "confidence": confidence,
            "target_language": target_language
        }

    elif fmt == "txt":
        if job_id:
            update_job_status(job_id, stage="generating", stage_label="Generating plain text file...")

        return {
            "bytes": translated_text.encode("utf-8"),
            "ext": "txt",
            "mime": "text/plain; charset=utf-8",
            "original_text": full_original_text,
            "translated_text": translated_text,
            "word_count": word_count,
            "detected_language": detected_lang,
            "confidence": confidence,
            "target_language": target_language
        }

    else:
        if job_id:
            update_job_status(job_id, stage="generating", stage_label="Generating translated PDF with Unicode fonts...")

        unicode_font = get_unicode_font_name()
        buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()
        story = []

        title_style = ParagraphStyle(
            'TitleStyle',
            fontName=unicode_font,
            fontSize=16,
            leading=22,
            textColor=colors.HexColor('#1E1B4B'),
            spaceAfter=12
        )

        body_style = ParagraphStyle(
            'BodyStyle',
            fontName=unicode_font,
            fontSize=10,
            leading=16,
            textColor=colors.HexColor('#334155'),
            spaceAfter=10
        )

        meta_style = ParagraphStyle(
            'MetaStyle',
            fontName=unicode_font,
            fontSize=9,
            leading=14,
            textColor=colors.HexColor('#64748B'),
            spaceAfter=16
        )

        story.append(Paragraph(f"DocFlow Translated Document ({target_language.capitalize()})", title_style))
        story.append(Paragraph(f"Detected Source: {detected_lang} ({confidence} confidence) | Target: {target_language.capitalize()} | Words: {word_count}", meta_style))
        story.append(Spacer(1, 10))

        for para in translated_text.split("\n\n"):
            clean_p = para.strip().replace("\n", " ")
            if clean_p:
                safe_p = clean_p.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                try:
                    story.append(Paragraph(safe_p, body_style))
                except Exception:
                    pass

        pdf_doc.build(story)
        out_bytes = buffer.getvalue()

        if job_id:
            update_job_status(job_id, stage="validating", stage_label="Validating output quality...")
        validate_output_pdf(out_bytes, min_pages=1)

        return {
            "bytes": out_bytes,
            "ext": "pdf",
            "mime": "application/pdf",
            "original_text": full_original_text,
            "translated_text": translated_text,
            "word_count": word_count,
            "detected_language": detected_lang,
            "confidence": confidence,
            "target_language": target_language
        }
