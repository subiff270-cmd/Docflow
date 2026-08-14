import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def _get_document_sections(doc_type: str, transcript: str) -> list[tuple[str, str]]:
    """Structures transcript text into appropriate sections based on document type."""
    now_str = datetime.now().strftime("%B %d, %Y")
    lines = [line.strip() for line in transcript.strip().splitlines() if line.strip()]
    full_text = "\n\n".join(lines) if lines else "No transcript content provided."

    dtype = doc_type.strip().lower()

    if "report" in dtype:
        return [
            ("EXECUTIVE SUMMARY", "This report synthesizes information recorded via the DocFlow Voice Workspace for review and archiving."),
            ("FINDINGS & CONTENT", full_text),
            ("CONCLUSION & RECOMMENDATIONS", "Based on the recorded findings, actionable steps and documentation should be maintained in compliance with standard operating procedures.")
        ]
    elif "assignment" in dtype:
        return [
            ("ACADEMIC OVERVIEW", f"Date of Submission: {now_str}\nSubject / Assignment Topic"),
            ("MAIN ANALYSIS & CONTENT", full_text),
            ("SUMMARY & FINDINGS", "The analysis above provides an overview of the investigated concepts and required deliverables.")
        ]
    elif "formal letter" in dtype or "letter" in dtype:
        return [
            ("DATE & RECIPIENT", f"Date: {now_str}\nTo Whom It May Concern\nDocFlow Verified Communications"),
            ("SUBJECT", f"Subject: Formal Communication - {lines[0][:60] if lines else 'Official Document'}"),
            ("BODY", full_text),
            ("CLOSING", "Sincerely,\nAuthorized Signatory\nDocFlow Document System")
        ]
    elif "email" in dtype:
        return [
            ("EMAIL HEADER", f"Date: {now_str}\nSubject: {lines[0][:60] if lines else 'Meeting / Discussion Notes'}"),
            ("MESSAGE", full_text),
            ("SIGNATURE", "Best regards,\nDocFlow Voice Workspace")
        ]
    elif "resume" in dtype:
        return [
            ("PROFESSIONAL SUMMARY", full_text),
            ("CORE COMPETENCIES", "• Strategic Planning & Execution\n• Communication & Documentation\n• Problem Resolution & Analytical Skills"),
            ("EXPERIENCE & ACHIEVEMENTS", "Detailed professional record based on the recorded session overview.")
        ]
    elif "meeting notes" in dtype or "meeting" in dtype:
        return [
            ("MEETING DETAILS", f"Date & Time: {now_str}\nStatus: Documented & Archived"),
            ("DISCUSSION & KEY POINTS", full_text),
            ("DECISIONS & ACTION ITEMS", "1. Review notes and distribute to stakeholders.\n2. Confirm next follow-up and implementation timeline.")
        ]
    elif "study notes" in dtype or "notes" in dtype:
        return [
            ("TOPIC OVERVIEW", f"Study Session: {now_str}"),
            ("CORE NOTES", full_text),
            ("KEY TAKEAWAYS", "Review the concepts outlined above for retention and exam/project preparation.")
        ]
    elif "project proposal" in dtype or "proposal" in dtype:
        return [
            ("EXECUTIVE SUMMARY", "This proposal details project objectives, methodology, and expected outcomes."),
            ("PROJECT SCOPE & METHODOLOGY", full_text),
            ("PROJECT TIMELINE & DELIVERABLES", "Milestones and delivery phases to be monitored according to project specification.")
        ]
    elif "article" in dtype:
        return [
            ("ARTICLE CONTENT", full_text),
            ("KEY TAKEAWAYS", "Summary insights and perspectives documented from the voice session.")
        ]
    elif "essay" in dtype:
        return [
            ("INTRODUCTION & THESIS", "The following exposition presents the central arguments and discussion."),
            ("MAIN BODY & ANALYSIS", full_text),
            ("CONCLUSION", "In summary, the documented perspectives align with the overarching theme.")
        ]
    else: # General Document
        return [
            ("DOCUMENT CONTENT", full_text)
        ]

def generate_voice_document(transcript_text: str, doc_type: str = "General Document", output_format: str = "docx") -> tuple[bytes, str]:
    """Generates real, professional DOCX, PDF, or TXT documents from transcript text."""
    clean_transcript = transcript_text.strip() if transcript_text else "No voice transcript recorded."
    sections = _get_document_sections(doc_type, clean_transcript)
    title_text = doc_type.upper()
    now_str = datetime.now().strftime("%B %d, %Y")

    # 1. Plain Text Output (.txt)
    if output_format == "txt":
        lines = [
            f"============================================================",
            f"DOCFLOW — {title_text}",
            f"Generated: {now_str}",
            f"============================================================\n"
        ]
        for sec_title, sec_content in sections:
            lines.append(f"--- {sec_title} ---")
            lines.append(sec_content + "\n")
        
        lines.append("------------------------------------------------------------")
        lines.append("Generated by DocFlow Voice to Document (https://docflow.com)")
        content = "\n".join(lines)
        return content.encode("utf-8"), "text/plain"

    # 2. PDF Output (.pdf) via ReportLab
    elif output_format == "pdf":
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'DocFlowTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#312e81"), # Indigo-900
            spaceAfter=4
        )
        
        meta_style = ParagraphStyle(
            'DocFlowMeta',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#64748b"), # Slate-500
            spaceAfter=14
        )
        
        sec_header_style = ParagraphStyle(
            'DocFlowSecHeader',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#4338ca"), # Indigo-700
            spaceBefore=12,
            spaceAfter=6
        )
        
        body_style = ParagraphStyle(
            'DocFlowBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=15,
            textColor=colors.HexColor("#1e293b"), # Slate-800
            spaceAfter=10
        )
        
        story = [
            Paragraph(title_text, title_style),
            Paragraph(f"DocFlow Voice Document • Created on {now_str}", meta_style),
            HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#4f46e5"), spaceAfter=14),
        ]
        
        for sec_title, sec_content in sections:
            story.append(Paragraph(sec_title, sec_header_style))
            for para in sec_content.splitlines():
                if para.strip():
                    story.append(Paragraph(para.strip(), body_style))
            story.append(Spacer(1, 6))
            
        story.append(Spacer(1, 14))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=8))
        story.append(Paragraph("Processed securely with DocFlow Document Suite (https://docflow.com)", meta_style))
        
        doc.build(story)
        return buffer.getvalue(), "application/pdf"

    # 3. Word Output (.docx) via python-docx
    else:
        doc = Document()
        
        # Document Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(title_text)
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(49, 46, 129) # Indigo 900
        
        # Meta info
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(f"DocFlow Voice Document • Generated on {now_str}")
        meta_run.font.name = 'Calibri'
        meta_run.font.size = Pt(9)
        meta_run.font.italic = True
        meta_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
        
        # Horizontal divider
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
        # Add Sections
        for sec_title, sec_content in sections:
            h = doc.add_heading(sec_title, level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            for r in h.runs:
                r.font.name = 'Calibri'
                r.font.bold = True
                r.font.color.rgb = RGBColor(67, 56, 202) # Indigo 700
                
            for para in sec_content.splitlines():
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    for r in p.runs:
                        r.font.name = 'Calibri'
                        r.font.size = Pt(11)
                        r.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        
        # Footer note
        doc.add_paragraph().paragraph_format.space_before = Pt(18)
        foot_p = doc.add_paragraph()
        foot_run = foot_p.add_run("Processed securely with DocFlow Document Suite (https://docflow.com)")
        foot_run.font.name = 'Calibri'
        foot_run.font.size = Pt(9)
        foot_run.font.italic = True
        foot_run.font.color.rgb = RGBColor(148, 163, 184) # Slate 400
        
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
